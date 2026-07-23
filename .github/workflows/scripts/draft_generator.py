import os
import json
from docx import Document

OUTPUT_DIR = "output"

def generate_draft(summary: dict, raw_txt_path: str):
    doc = Document()
    
    # 標題
    doc.add_heading("刑事文件草稿 (AI輔助整理)", level=1)
    
    # 基本案件資訊
    doc.add_heading("案件資訊", level=2)
    doc.add_paragraph(f"案號：{summary.get('case_number', '')}")
    doc.add_paragraph(f"法院：{summary.get('court', '')}")
    doc.add_paragraph(f"資料來源：{summary.get('source', '')}")
    
    # 待確認事項
    doc.add_heading("待確認事項", level=2)
    need_verify = summary.get("need_verify", [])
    if need_verify:
        for item in need_verify:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("（尚未填寫，請手動補充）")
    
    # 證據整理（法院認定部分）
    doc.add_heading("法院認定摘要", level=2)
    findings = summary.get("court_finding", [])
    if findings:
        for item in findings:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("（尚未從判決書整理，請自行填入）")
    
    # 你的主張（手動區域）
    doc.add_heading("當事人主張", level=2)
    claims = summary.get("user_claim", [])
    if claims:
        for item in claims:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("（請自行填寫你的說法）")
    
    # 地址欄位（留白）
    doc.add_heading("送達地址", level=2)
    doc.add_paragraph("送達地址：_______________________________________________")
    doc.add_paragraph("送達代收人：__________________  電話：__________________")
    
    # AI 提醒與免責
    doc.add_heading("AI 提醒", level=2)
    doc.add_paragraph("⚡ 本文件為AI輔助整理草稿，需本人確認，不代表法院正式文件。")
    doc.add_paragraph("⚡ 請務必核對判決書原文，並於上訴期限內向法院遞交正式狀紙。")
    
    # 附上原始判決書節錄（可選）
    doc.add_heading("原始判決書節錄", level=2)
    try:
        with open(raw_txt_path, "r", encoding="utf-8") as f:
            raw_text = f.read()
        # 只取前2000字避免文件過大
        doc.add_paragraph(raw_text[:2000])
    except:
        doc.add_paragraph("（無法讀取原始檔案）")
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(os.path.join(OUTPUT_DIR, "appeal_draft.docx"))
