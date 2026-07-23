import os
import re
from datetime import datetime
from docx import Document

RAW_FILE = "cases/115簡238_raw.txt"
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 自動產生檔名：上訴理由狀_115年度簡字第238號_1150724.docx
today = datetime.now().strftime("%Y%m%d")  # 1150724
case_no = "115年度簡字第238號"
filename = f"上訴理由狀_{case_no}_{today}.docx"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, filename)

YOUR_CLAIMS = {
    "case_number": case_no,
    "court": "臺灣臺中地方法院",
    "name": "徐志曆",
    "arguments": [
        {
            "title": "原判決所憑「自白」存有重大疑義，事實認定基礎已生動搖",
            "body": [
                "原判決以「被告於準備程序時之自白」為主要論罪依據。然上訴人於準備程序所述真意，係表明「不應與小朋友互動造成誤會」，意在澄清主觀上無騷擾故意，從未承認曾以手觸摸他人臀部。",
                "筆錄記載與上訴人真意顯有出入，該「自白」是否成立、內容為何，實有再行勘驗準備程序錄音及筆錄之必要。"
            ]
        },
        {
            "title": "原審僅憑監視器「翻拍照片」即為有罪推論，應重新調查完整連續影像",
            "body": [
                "原審引用「現場監視錄影畫面翻拍照片」為補強證據，然靜態截圖無法重現動態歷程，更難以單一畫面斷言身體接觸已然發生。",
                "上訴人始終主張手部當時位在自己座椅附近，並無碰觸他人臀部之舉。請求調取美芝城早餐店114年8月30日中午12時前後之原始監視器錄影檔案，並當庭勘驗。"
            ]
        },
        {
            "title": "上訴人主觀上毫無性騷擾之犯意，原審量刑未充分審酌",
            "body": [
                "上訴人當時所處位置、與在場人之距離、舉動之目的，均與蓄意侵犯他人身體隱私之情形有別。上訴人因見現場小朋友而與之互動，係出於善意關心，絕無破壞他人身體界線之意圖。"
            ]
        },
        {
            "title": "執行程序及住宅進入合法性之調查聲請",
            "body": [
                "上訴人主張，於本案相關執行過程中，執法人員進入住處並限制本人行動之經過，有程序合法性疑義。",
                "上訴人於當時因認執法人員未提示搜索票，且未充分說明進入住宅之法律依據，故對於執行行為有所疑問並產生抗拒。",
                "爰請求法院調取當日執行文件、勤務紀錄、密錄器影像及其他可還原現場經過之資料，以確認執行程序是否符合刑事訴訟法及憲法保障之正當法律程序。"
            ]
        },
        {
            "title": "原審未審酌對上訴人有利之量刑因子，應再予從輕",
            "body": [
                "上訴人目前有正當工作，照顧年邁家人，生活穩定。對於本案造成告訴人不快，始終願表達歉意，並非毫無悔意。",
                "原審量刑時，就上開對上訴人有利之因子未見具體審認，所處刑度顯然過重，有再予減輕之必要。"
            ]
        }
    ],
    "need_verify": [
        "準備程序錄音及筆錄",
        "原始監視器錄影檔案",
        "當日執行程序相關文件",
        "勤務紀錄及密錄器影像",
        "判決正本實際送達日期"
    ]
}

def generate_draft():
    raw_text = ""
    if os.path.exists(RAW_FILE):
        with open(RAW_FILE, "r", encoding="utf-8") as f:
            raw_text = f.read()

    court = YOUR_CLAIMS["court"]
    m = re.search(r"臺灣\w+地方法院", raw_text)
    if m:
        court = m.group(0)

    doc = Document()
    doc.add_heading("刑事上訴理由狀（AI輔助草稿）", level=1)

    doc.add_paragraph(f"案號：{YOUR_CLAIMS['case_number']}")
    doc.add_paragraph(f"原審法院：{court}")
    doc.add_paragraph(f"上訴人：{YOUR_CLAIMS['name']}")
    doc.add_paragraph(f"日期：中華民國 115 年 {datetime.now().strftime('%m')} 月 {datetime.now().strftime('%d')} 日")

    doc.add_heading("上訴理由", level=2)
    for i, arg in enumerate(YOUR_CLAIMS["arguments"], 1):
        doc.add_heading(f"{i}. {arg['title']}", level=3)
        for line in arg["body"]:
            doc.add_paragraph(line)

    doc.add_heading("待法院調查之證據及事項", level=2)
    for item in YOUR_CLAIMS["need_verify"]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("送達地址", level=2)
    doc.add_paragraph("送達地址：_______________________________________________")
    doc.add_paragraph("送達代收人：__________________  電話：__________________")

    doc.add_heading("注意", level=2)
    doc.add_paragraph("本文件由 AI 根據你提供之主張自動產生，僅供草稿使用。")
    doc.add_paragraph("遞交法院前，務必由你本人確認全部內容並親筆簽名或蓋章。")
    doc.add_paragraph("上訴期限以你實際收受判決正本翌日起算20日，請速向臺中地院刑事科查詢送達日期。")

    doc.add_heading("具狀人", level=2)
    doc.add_paragraph(f"上訴人：{YOUR_CLAIMS['name']}（簽章）")

    doc.save(OUTPUT_FILE)
    print(f"✅ 已產生 {OUTPUT_FILE}")
    return OUTPUT_FILE

if __name__ == "__main__":
    generate_draft()
